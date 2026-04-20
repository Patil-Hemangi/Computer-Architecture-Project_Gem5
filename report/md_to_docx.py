"""
Convert EEL6764_HNSW_Report.md to a formatted Word document.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH  = r"d:\GEM5_Class_Project\report\EEL6764_HNSW_Report.md"
OUT_PATH = r"d:\GEM5_Class_Project\report\EEL6764_HNSW_Report_final6.docx"

def set_heading_color(paragraph, r, g, b):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)

def add_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '999999')
        tblBorders.append(el)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def apply_inline(run_obj, text):
    """Apply bold/italic/code inline formatting to a run."""
    run_obj.text = text

def add_paragraph_with_inline(doc, text, style=None, indent=None):
    """Add a paragraph parsing **bold**, *italic*, and `code` inline."""
    p = doc.add_paragraph(style=style)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    # Split on bold/italic/code markers
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            r = p.add_run(part[3:-3])
            r.bold = True; r.italic = True
        elif part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = p.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        else:
            p.add_run(part)
    return p

def parse_table(lines):
    """Parse markdown table lines into list of row lists."""
    rows = []
    for line in lines:
        if re.match(r'\|[-| :]+\|', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    return rows

def clean_md(text):
    """Strip markdown bold/italic/code markers from plain text."""
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\$\$.*?\$\$', '[formula]', text, flags=re.DOTALL)
    text = re.sub(r'\$.*?\$', '[formula]', text)
    return text.strip()

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(MD_PATH, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code  = False
    code_buf = []

    while i < len(lines):
        raw = lines[i].rstrip('\n')
        stripped = raw.strip()

        # ── fenced code block ──────────────────────────────────────────────
        if stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                for cl in code_buf:
                    run = p.add_run(cl + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # ── horizontal rule ────────────────────────────────────────────────
        if re.match(r'^-{3,}$', stripped):
            doc.add_paragraph()
            i += 1
            continue

        # ── image lines ────────────────────────────────────────────────────
        if stripped.startswith('!['):
            m_img = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
            if m_img:
                alt_text = m_img.group(1)
                img_rel  = m_img.group(2)
                # resolve relative to report dir
                import os
                report_dir = os.path.dirname(MD_PATH)
                img_path = os.path.normpath(os.path.join(report_dir, img_rel))
                if os.path.exists(img_path):
                    doc.add_paragraph()
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(img_path, width=Inches(5.5))
                    # caption
                    cap = doc.add_paragraph(alt_text)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap.runs:
                        run.italic = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x55,0x55,0x55)
                    doc.add_paragraph()
                else:
                    p = doc.add_paragraph(f'[Figure: {alt_text}]')
                    for run in p.runs:
                        run.italic = True
                        run.font.color.rgb = RGBColor(0x99,0x99,0x99)
            i += 1
            continue

        # ── markdown table ─────────────────────────────────────────────────
        if stripped.startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i].strip())
                i += 1
            rows = parse_table(tbl_lines)
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            # normalise rows
            rows = [r + [''] * (ncols - len(r)) for r in rows]
            t = doc.add_table(rows=len(rows), cols=ncols)
            t.style = 'Table Grid'
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = t.cell(ri, ci)
                    cell.text = ''
                    p2 = cell.paragraphs[0]
                    plain = clean_md(cell_text)
                    run = p2.add_run(plain)
                    run.font.size = Pt(9)
                    if ri == 0:
                        run.bold = True
                        cell._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'E8EFF6')
                        cell._tc.tcPr.append(shd)
            add_table_border(t)
            doc.add_paragraph()
            continue

        # ── headings ───────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            title = clean_md(m.group(2))
            hmap  = {1:'Heading 1', 2:'Heading 2', 3:'Heading 3', 4:'Heading 4'}
            h = doc.add_heading(title, level=min(level, 4))
            h.style = doc.styles[hmap.get(level, 'Heading 4')]
            # colour headings
            colors = {1:(0x1B,0x4F,0x72), 2:(0x1A,0x5,0x76), 3:(0x21,0x61,0x8A), 4:(0x2E,0x86,0xC1)}
            set_heading_color(h, *colors.get(level, (0,0,0)))
            i += 1
            continue

        # ── blockquote ─────────────────────────────────────────────────────
        if stripped.startswith('>'):
            text = re.sub(r'^>\s*', '', stripped)
            p = add_paragraph_with_inline(doc, clean_md(text))
            p.paragraph_format.left_indent  = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55,0x55,0x55)
            i += 1
            continue

        # ── bullet list ────────────────────────────────────────────────────
        if re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            add_paragraph_with_inline(doc, clean_md(text), style='List Bullet')
            i += 1
            continue

        # ── numbered list ──────────────────────────────────────────────────
        if re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            add_paragraph_with_inline(doc, clean_md(text), style='List Number')
            i += 1
            continue

        # ── blank line ─────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── normal paragraph ───────────────────────────────────────────────
        add_paragraph_with_inline(doc, clean_md(stripped))
        i += 1

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")

if __name__ == '__main__':
    main()
